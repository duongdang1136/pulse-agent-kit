from pathlib import Path
import json

from openpyxl import Workbook
from docx import Document

from pulse.cli import main
from pulse.knowledge import import_knowledge, list_knowledge
from pulse.project import create_project
from pulse.rag.pipeline import build_index


def make_repo(tmp_path: Path) -> Path:
    for folder in ("agents", "knowledge", "scripts"):
        (tmp_path / folder).mkdir()
    (tmp_path / "scripts" / "rag.py").write_text("print('stub')\n", encoding="utf-8")
    return tmp_path


def test_import_markdown_and_text(tmp_path: Path) -> None:
    root = make_repo(tmp_path)
    create_project(root, "FPTPlay")
    source = tmp_path / "incoming"
    source.mkdir()
    (source / "Login Flow.md").write_text("# Login\n\nGuest cannot follow teams.", encoding="utf-8")
    (source / "notes.txt").write_text("Live Activity notes", encoding="utf-8")

    results = import_knowledge(root, "fptplay", source)

    assert len(results) == 2
    assert all(item.status == "created" for item in results)
    pages = list_knowledge(root, "fptplay")
    assert len(pages) == 2
    login = (root / "knowledge/projects/fptplay/pages/login-flow.md").read_text(encoding="utf-8")
    assert 'source_type: "md"' in login
    assert "Guest cannot follow teams" in login
    assert (root / "projects/fptplay/source-docs/Login Flow.md").exists()


def test_import_docx_and_xlsx(tmp_path: Path) -> None:
    root = make_repo(tmp_path)
    create_project(root, "FPTPlay")

    docx_path = tmp_path / "requirements.docx"
    document = Document()
    document.add_heading("Live Activity", level=1)
    document.add_paragraph("Show match state on lock screen.")
    document.save(docx_path)

    xlsx_path = tmp_path / "rules.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Business Rules"
    sheet.append(["ID", "Rule"])
    sheet.append(["BR-01", "Guest cannot follow a team"])
    workbook.save(xlsx_path)

    assert import_knowledge(root, "fptplay", docx_path)[0].status == "created"
    assert import_knowledge(root, "fptplay", xlsx_path)[0].status == "created"
    assert "Show match state" in (root / "knowledge/projects/fptplay/pages/requirements.md").read_text(encoding="utf-8")
    assert "BR-01" in (root / "knowledge/projects/fptplay/pages/rules.md").read_text(encoding="utf-8")


def test_cli_knowledge_import(tmp_path: Path, monkeypatch) -> None:
    root = make_repo(tmp_path)
    create_project(root, "FPTPlay")
    source = tmp_path / "brief.txt"
    source.write_text("Feature brief", encoding="utf-8")
    monkeypatch.chdir(root)

    assert main(["knowledge", "import", "fptplay", str(source)]) == 0
    assert main(["knowledge", "list", "fptplay"]) == 0
    assert main(["doctor"]) == 0


def test_cli_knowledge_import_accepts_category(tmp_path: Path, monkeypatch) -> None:
    root = make_repo(tmp_path)
    create_project(root, "FPTPlay")
    source = tmp_path / "architecture.txt"
    source.write_text("API gateway routes requests.", encoding="utf-8")
    monkeypatch.chdir(root)

    assert main(["knowledge", "import", "fptplay", str(source), "--category", "Architecture"]) == 0

    page = root / "knowledge/projects/fptplay/pages/architecture.md"
    assert 'category: "Architecture"' in page.read_text(encoding="utf-8")


def test_import_skips_unchanged_source_without_touching_manifest(tmp_path: Path) -> None:
    root = make_repo(tmp_path)
    create_project(root, "FPTPlay")
    source = tmp_path / "brief.txt"
    source.write_text("Feature brief", encoding="utf-8")

    first = import_knowledge(root, "fptplay", source)[0]
    assert first.status == "created"

    manifest_path = root / "knowledge/projects/fptplay/manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["documents"][0]["index_fingerprint"] = "fingerprint-v1"
    manifest_path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")

    second = import_knowledge(root, "fptplay", source)[0]

    assert second.status == "skipped"
    reloaded = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert reloaded["documents"][0]["index_fingerprint"] == "fingerprint-v1"


def test_import_without_overwrite_does_not_update_manifest_for_changed_source(tmp_path: Path) -> None:
    root = make_repo(tmp_path)
    create_project(root, "FPTPlay")
    source = tmp_path / "brief.txt"
    source.write_text("Original brief", encoding="utf-8")

    first = import_knowledge(root, "fptplay", source)[0]
    original_page = first.page_path.read_text(encoding="utf-8")
    original_manifest = json.loads((root / "knowledge/projects/fptplay/manifest.json").read_text(encoding="utf-8"))
    original_checksum = original_manifest["documents"][0]["checksum"]

    source.write_text("Changed brief", encoding="utf-8")
    second = import_knowledge(root, "fptplay", source)[0]

    assert second.status == "skipped"
    assert first.page_path.read_text(encoding="utf-8") == original_page
    reloaded = json.loads((root / "knowledge/projects/fptplay/manifest.json").read_text(encoding="utf-8"))
    assert reloaded["documents"][0]["checksum"] == original_checksum


def test_import_with_overwrite_updates_existing_page_and_resets_index(tmp_path: Path) -> None:
    root = make_repo(tmp_path)
    create_project(root, "FPTPlay")
    source = tmp_path / "brief.txt"
    source.write_text("Original brief", encoding="utf-8")

    first = import_knowledge(root, "fptplay", source)[0]
    manifest_path = root / "knowledge/projects/fptplay/manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["documents"][0]["index_fingerprint"] = "fingerprint-v1"
    manifest_path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")

    source.write_text("Changed brief", encoding="utf-8")
    second = import_knowledge(root, "fptplay", source, overwrite=True)[0]

    assert second.status == "updated"
    assert second.document_id == first.document_id
    assert "Changed brief" in second.page_path.read_text(encoding="utf-8")
    reloaded = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert reloaded["documents"][0]["index_fingerprint"] == ""


def test_cli_rag_query_can_include_shared_knowledge(tmp_path: Path, monkeypatch, capsys) -> None:
    root = make_repo(tmp_path)
    create_project(root, "FPTPlay")
    project_root = root / "knowledge/projects/fptplay"
    shared_root = root / "knowledge/shared"
    (shared_root / "pages").mkdir(parents=True)

    (project_root / "pages" / "notification.md").write_text(
        "---\nid: notification\n---\n\nFPTPlay notification scheduling uses campaign windows.",
        encoding="utf-8",
    )
    (project_root / "manifest.json").write_text(json.dumps({
        "schema_version": 2,
        "project": "fptplay",
        "documents": [{
            "id": "notification",
            "title": "Notification",
            "project": "fptplay",
            "source": "notification.txt",
            "page": "pages/notification.md",
            "checksum": "",
            "source_type": "txt",
            "status": "imported",
            "created_at": "2026-01-01",
            "updated_at": "2026-01-01",
            "index_fingerprint": "",
        }],
    }), encoding="utf-8")
    (shared_root / "pages" / "push.md").write_text(
        "---\nid: push\n---\n\nPush notification scheduling should consider timezone and quiet hours.",
        encoding="utf-8",
    )
    (shared_root / "manifest.json").write_text(json.dumps({
        "schema_version": 2,
        "project": "shared",
        "documents": [{
            "id": "push",
            "title": "Push",
            "project": "shared",
            "source": "push.txt",
            "page": "pages/push.md",
            "checksum": "",
            "source_type": "txt",
            "status": "imported",
            "created_at": "2026-01-01",
            "updated_at": "2026-01-01",
            "index_fingerprint": "",
        }],
    }), encoding="utf-8")
    build_index(project_root, "fptplay", provider="hash", force=True)
    build_index(shared_root, "shared", provider="hash", force=True)
    monkeypatch.chdir(root)

    assert main(["rag", "query", "fptplay", "notification scheduling", "--include-shared", "--top-k", "4"]) == 0

    output = capsys.readouterr().out
    assert "scope=project" in output
    assert "scope=shared" in output
