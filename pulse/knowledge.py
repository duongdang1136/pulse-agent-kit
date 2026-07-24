from __future__ import annotations

import hashlib
import json
import re
import shutil
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable

from pulse.project import ProjectInfo, read_project, slugify

SUPPORTED_EXTENSIONS = {".md", ".markdown", ".txt", ".pdf", ".docx", ".xlsx"}


@dataclass(frozen=True)
class ImportResult:
    source: Path
    copied_to: Path
    page_path: Path
    document_id: str
    status: str


def import_knowledge(
    root: Path,
    project_identifier: str,
    source: Path,
    *,
    copy_source: bool = True,
    overwrite: bool = False,
) -> list[ImportResult]:
    project = read_project(root, project_identifier)
    source = source.expanduser().resolve()
    if not source.exists():
        raise FileNotFoundError(f"Source path was not found: {source}")

    files = list(_iter_supported_files(source))
    if not files:
        raise ValueError("No supported files found. Supported: md, txt, pdf, docx, xlsx")

    results: list[ImportResult] = []
    used_ids: set[str] = set()
    for file_path in files:
        relative = file_path.name if source.is_file() else file_path.relative_to(source).as_posix()
        document_id = _unique_document_id(relative, used_ids)
        used_ids.add(document_id)
        results.append(
            _import_file(root, project, file_path, relative, document_id, copy_source, overwrite)
        )
    _write_manifest(root, project, results)
    return results


def _iter_supported_files(source: Path) -> Iterable[Path]:
    if source.is_file():
        if source.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield source
        return
    for path in sorted(source.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS and not _is_hidden(path, source):
            yield path


def _is_hidden(path: Path, base: Path) -> bool:
    try:
        parts = path.relative_to(base).parts
    except ValueError:
        parts = path.parts
    return any(part.startswith(".") for part in parts)


def _import_file(
    root: Path,
    project: ProjectInfo,
    source: Path,
    relative: str,
    document_id: str,
    copy_source: bool,
    overwrite: bool,
) -> ImportResult:
    workspace_source = root / project.workspace_path / "source-docs" / relative
    page = root / project.knowledge_path / "pages" / f"{document_id}.md"
    if page.exists() and not overwrite:
        return ImportResult(source, workspace_source, page, document_id, "skipped")

    text = extract_text(source)
    if not text.strip():
        raise ValueError(f"No readable text found in {source}")

    if copy_source:
        workspace_source.parent.mkdir(parents=True, exist_ok=True)
        if source.resolve() != workspace_source.resolve():
            shutil.copy2(source, workspace_source)

    page.parent.mkdir(parents=True, exist_ok=True)
    page.write_text(
        _render_page(project, source, relative, document_id, text), encoding="utf-8"
    )
    return ImportResult(source, workspace_source, page, document_id, "imported")


def extract_text(path: Path) -> str:
    extension = path.suffix.lower()
    if extension in {".md", ".markdown", ".txt"}:
        return path.read_text(encoding="utf-8-sig", errors="replace")
    if extension == ".pdf":
        return _extract_pdf(path)
    if extension == ".docx":
        return _extract_docx(path)
    if extension == ".xlsx":
        return _extract_xlsx(path)
    raise ValueError(f"Unsupported file type: {extension}")


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF import requires pypdf. Run: python -m pip install -e .") from exc
    reader = PdfReader(str(path))
    sections = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            sections.append(f"## Page {index}\n\n{text.strip()}")
    return "\n\n".join(sections)


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX import requires python-docx. Run: python -m pip install -e .") from exc
    document = Document(str(path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text.strip())
    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if rows:
            blocks.append(_markdown_table(rows, f"Table {table_index}"))
    return "\n\n".join(blocks)


def _extract_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("XLSX import requires openpyxl. Run: python -m pip install -e .") from exc
    workbook = load_workbook(path, read_only=True, data_only=True)
    sections: list[str] = []
    for worksheet in workbook.worksheets:
        rows = []
        for row in worksheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                rows.append(values)
        if rows:
            sections.append(_markdown_table(rows, f"Sheet: {worksheet.title}"))
    workbook.close()
    return "\n\n".join(sections)


def _markdown_table(rows: list[list[str]], title: str) -> str:
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    body = normalized[1:]
    lines = [f"## {title}", "", "| " + " | ".join(_escape_cell(v) for v in header) + " |"]
    lines.append("| " + " | ".join("---" for _ in range(width)) + " |")
    lines.extend("| " + " | ".join(_escape_cell(v) for v in row) + " |" for row in body)
    return "\n".join(lines)


def _escape_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\r", " ").replace("\n", " ")


def _render_page(
    project: ProjectInfo, source: Path, relative: str, document_id: str, text: str
) -> str:
    title = source.stem.replace("-", " ").replace("_", " ").strip().title()
    checksum = hashlib.sha256(source.read_bytes()).hexdigest()
    normalized = _normalize_text(text)
    return (
        "---\n"
        f'id: "{document_id}"\n'
        f'title: "{_yaml_escape(title)}"\n'
        'type: "knowledge"\n'
        'category: "Product"\n'
        f'project: "{_yaml_escape(project.slug)}"\n'
        f'source_file: "{_yaml_escape(relative)}"\n'
        f'source_type: "{source.suffix.lower().lstrip(".")}"\n'
        f'updated_at: "{date.today().isoformat()}"\n'
        f'checksum_sha256: "{checksum}"\n'
        "tags:\n"
        f'  - "{_yaml_escape(project.slug)}"\n'
        "---\n\n"
        f"# {title}\n\n"
        f"> Imported from `{relative}`. Treat this page as project evidence, not an automatically verified decision.\n\n"
        f"{normalized}\n"
    )


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def _yaml_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace('"', '\\"')


def _unique_document_id(relative: str, used_ids: set[str]) -> str:
    stem = Path(relative).with_suffix("").as_posix().replace("/", "-")
    base = slugify(stem)
    candidate = base
    index = 2
    while candidate in used_ids:
        candidate = f"{base}-{index}"
        index += 1
    return candidate


def _write_manifest(root: Path, project: ProjectInfo, results: list[ImportResult]) -> None:
    from pulse.rag.manifest import load_manifest, save_manifest

    manifest_path = root / project.knowledge_path / "manifest.json"
    payload = load_manifest(manifest_path, project.slug)
    existing = {item.get("id"): item for item in payload["documents"] if item.get("id")}
    today = date.today().isoformat()
    for result in results:
        checksum = hashlib.sha256(result.source.read_bytes()).hexdigest()
        current = existing.get(result.document_id, {})
        current.update({
            "id": result.document_id,
            "title": result.source.stem.replace("-", " ").replace("_", " ").strip().title(),
            "project": project.slug,
            "source": str(result.source),
            "page": result.page_path.relative_to(root / project.knowledge_path).as_posix(),
            "checksum": checksum,
            "source_type": result.source.suffix.lower().lstrip("."),
            "status": "imported" if result.status == "imported" else current.get("status", "imported"),
            "created_at": current.get("created_at", today),
            "updated_at": today,
            "metadata": {"copied_to": result.copied_to.as_posix()},
        })
        if result.status == "imported":
            current["index_fingerprint"] = ""
        existing[result.document_id] = current
    payload["schema_version"] = 2
    payload["project"] = project.slug
    payload["documents"] = sorted(existing.values(), key=lambda item: item["id"])
    save_manifest(manifest_path, payload)


def list_knowledge(root: Path, project_identifier: str) -> list[Path]:
    project = read_project(root, project_identifier)
    pages = root / project.knowledge_path / "pages"
    return sorted(path for path in pages.glob("*.md") if path.name.lower() != "readme.md")
